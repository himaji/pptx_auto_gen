import tkinter as tk
import tkinter.ttk as ttk
import sells_report_auto_gen
import pandas as pd
import datetime as dt
from functools import partial
import openpyxl
import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm

class MainWindow(tk.Frame):
    def __init__(self, master=None, parent=None):
        super().__init__(master)
        self.master = master
        self.master.geometry("800x700")
        self.master.title("売上報告書できるやつ")
        self.master.grid_rowconfigure(0, weight=1)
        self.master.grid_columnconfigure(0, weight=1)

        self.set_data()
        self.create_widgets()

    def set_data(self):
        self.df_sales_management = pd.read_excel("../output/sales_management.xlsx", sheet_name="販売個数", header=0, index_col=None)
        self.colname_list = ["日付", "弁当名等", "搬入個数"]  # 結果に表示させる列名
        self.width_list = [100, 200, 50]
        self.search_col = "日付"  # 検索キーワードの入力されている列名

        self.df_menu_master = pd.read_excel("../output/sales_management.xlsx", sheet_name="弁当名マスタ", header=0, index_col=None)

    def create_widgets(self):
        self.pw_main = ttk.PanedWindow(self.master, orient="vertical")
        self.pw_main.pack(expand=True, fill=tk.BOTH, side="left")
        self.pw_top = ttk.PanedWindow(self.pw_main, orient="horizontal", height=30)
        self.pw_main.add(self.pw_top)
        self.pw_bottom = ttk.PanedWindow(self.pw_main, orient="vertical")
        self.pw_main.add(self.pw_bottom)
        self.pw_bottom2 = ttk.PanedWindow(self.pw_main, orient="vertical")
        self.pw_main.add(self.pw_bottom2)

        self.create_input_frame(self.pw_top)
        self.create_tree(self.pw_bottom)
        self.create_update_frame(self.pw_bottom2)

    def create_input_frame(self, parent):
        fm_input = ttk.Frame(parent, )
        parent.add(fm_input)
        lbl_keyword = ttk.Label(fm_input, text="キーワード", width=7)
        lbl_keyword.grid(row=1, column=1, padx=2, pady=2)
        self.keyword = tk.StringVar()
        ent_keyword = ttk.Entry(fm_input, justify="left", textvariable=self.keyword)
        ent_keyword.grid(row=1, column=2, padx=2, pady=2)
        ent_keyword.bind("<Return>", self.search)

    def create_tree(self, parent):
        self.result_text = tk.StringVar()
        lbl_result = ttk.Label(parent, textvariable=self.result_text)
        parent.add(lbl_result)
        self.tree = ttk.Treeview(parent)
        self.tree["column"] = self.colname_list
        self.tree["show"] = "headings"
        self.tree.bind("<Double-1>", self.onDouble)
        for i, (colname, width) in enumerate(zip(self.colname_list, self.width_list)):
            self.tree.heading(i, text=colname)
            self.tree.column(i, width=width)
        parent.add(self.tree)

    def search(self, event=None):
        keyword = dt.datetime.strptime(self.keyword.get(), '%Y-%m-%d')
        result = self.df_sales_management[self.df_sales_management['日付'] == keyword]
        self.update_tree_by_search_result(result)

    def create_update_frame(self, parent):
        fm_update = ttk.Frame(parent,)
        parent.add(fm_update)
        lbl_update1 = ttk.Label(fm_update, text="1")
        lbl_update1.grid(row=1, column=1, padx=2, pady=2)
        self.ent_update1 = ttk.Entry(fm_update, justify="left")
        self.ent_update1.grid(row=1, column=2, padx=2, pady=2)
        lbl_update2 = ttk.Label(fm_update, text="2")
        lbl_update2.grid(row=2, column=1, padx=2, pady=2)
        self.ent_update2 = ttk.Entry(fm_update, justify="left")
        self.ent_update2.grid(row=2, column=2, padx=2, pady=2)
        lbl_update3 = ttk.Label(fm_update, text="3")
        lbl_update3.grid(row=3, column=1, padx=2, pady=2)
        self.ent_update3 = ttk.Entry(fm_update, justify="left")
        self.ent_update3.grid(row=3, column=2, padx=2, pady=2)
        lbl_update4 = ttk.Label(fm_update, text="4")
        lbl_update4.grid(row=4, column=1, padx=2, pady=2)
        self.ent_update4 = ttk.Entry(fm_update, justify="left")
        self.ent_update4.grid(row=4, column=2, padx=2, pady=2)
        lbl_update5 = ttk.Label(fm_update, text="5")
        lbl_update5.grid(row=5, column=1, padx=2, pady=2)
        self.ent_update5 = ttk.Entry(fm_update, justify="left")
        self.ent_update5.grid(row=5, column=2, padx=2, pady=2)
        lbl_update6 = ttk.Label(fm_update, text="6")
        lbl_update6.grid(row=6, column=1, padx=2, pady=2)
        self.ent_update6 = ttk.Entry(fm_update, justify="left")
        self.ent_update6.grid(row=6, column=2, padx=2, pady=2)
        lbl_update7 = ttk.Label(fm_update, text="7")
        lbl_update7.grid(row=7, column=1, padx=2, pady=2)
        self.ent_update7 = ttk.Entry(fm_update, justify="left")
        self.ent_update7.grid(row=7, column=2, padx=2, pady=2)
        lbl_update8 = ttk.Label(fm_update, text="8")
        lbl_update8.grid(row=8, column=1, padx=2, pady=2)
        self.ent_update8 = ttk.Entry(fm_update, justify="left")
        self.ent_update8.grid(row=8, column=2, padx=2, pady=2)
        lbl_update9 = ttk.Label(fm_update, text="9")
        lbl_update9.grid(row=9, column=1, padx=2, pady=2)
        self.ent_update9 = ttk.Entry(fm_update, justify="left")
        self.ent_update9.grid(row=9, column=2, padx=2, pady=2)
        lbl_update10 = ttk.Label(fm_update, text="10")
        lbl_update10.grid(row=10, column=1, padx=2, pady=2)
        self.ent_update10 = ttk.Entry(fm_update, justify="left")
        self.ent_update10.grid(row=10, column=2, padx=2, pady=2)
        lbl_update_coupon = ttk.Label(fm_update, text="クーポン利用")
        lbl_update_coupon.grid(row=11, column=1, padx=2, pady=2)
        self.ent_update_coupon = ttk.Entry(fm_update, justify="left")
        self.ent_update_coupon.grid(row=11, column=2, padx=2, pady=2)

        btn_current = ttk.Button(fm_update, text="更新", command=lambda: self.update_result())
        btn_current.grid(row=12, column=1, padx=2, pady=2)

        btn_current = ttk.Button(fm_update, text="売上報告作成", command=lambda: self.docx_auto_gen())
        btn_current.grid(row=12, column=2, padx=2, pady=2)


    def docx_auto_gen(self):
        date = dt.datetime.strptime(self.keyword.get(), '%Y-%m-%d')
        colname_list = ["日付", "販売価格", "弁当名等", "搬入個数", "販売個数"]  # 結果に表示させる列名
        width_list = [100, 200, 50]
        search_col = "日付"  # 検索キーワードの入力されている列名

        df_menu_master = pd.read_excel("../output/sales_management.xlsx", sheet_name="弁当名マスタ", header=0, index_col=None)

        self.df_sales_management.loc[self.df_sales_management['弁当名等']=="50円引きクーポン利用", ['販売価格']] =  self.df_sales_management['割引等単価']
        self.df_sales_management.loc[self.df_sales_management['弁当名等']=="50円引きクーポン利用", ['販売個数']] =  self.df_sales_management['クーポン利用']

        df_formated = self.df_sales_management[colname_list]
        df_formated = df_formated[df_formated["日付"] == date]
        df_formated = df_formated.fillna(0)

        array = []
        for value, sales_quantity in zip(df_formated["販売価格"], df_formated["販売個数"]):
            if value != None and sales_quantity != None:
                array.append(int(value*sales_quantity))

        s = pd.Series(array,  index=df_formated.index, name='合計')
        df_formated = pd.concat([df_formated, s], axis=1)

        document = Document()
        document.add_heading('売上日報', 0)

        p = document.add_paragraph(str(date.date()))
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p = document.add_paragraph('作成者  横田 秀喜')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        table = document.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        table.add_column(Mm(30))
        table.add_column(Mm(30))
        table.add_column(Mm(40))
        table.add_column(Mm(30))
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '商品名'
        # hdr_cells[0].width = Inches(0.2)
        hdr_cells[1].text = '単価'
        hdr_cells[2].text = '数量'
        hdr_cells[3].text = '搬入個数'
        hdr_cells[4].text = '合計'
        for menu_name, value, delivered_quantity, saled_quantity, total in zip(df_formated["弁当名等"],df_formated["販売価格"],df_formated["販売個数"],df_formated["搬入個数"],df_formated["合計"]):
            row_cells = table.add_row().cells
            row_cells[0].text = menu_name
            # row_cells[1].text = str("{:,}".format(int(value)))
            # row_cells[1].text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p = document.add_paragraph(str("{:,}".format(int(value))))
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[1].paragraphs[0] = p
            # row_cells[1].add_paragraph(str("{:,}".format(int(value))))
            # row_cells[1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # row_cells[1].text = p
            row_cells[2].text = str("{:,}".format(int(delivered_quantity)))
            row_cells[3].text = str("{:,}".format(int(saled_quantity)))
            row_cells[4].text = str("{:,}".format(int(total)))
        row_cells = table.add_row().cells
        row_cells[4].text = str("{:,}".format(int(df_formated["合計"].sum())))
        document.save('../output/demo.docx') 

    def update_result(self):
        if self.keyword.get() != "" :
            print("update_result")
            self.ent_index = []
            if self.ent_update1.get() != "":
                self.ent_index.append(self.ent_update1.get())
            if self.ent_update2.get() != "":
                self.ent_index.append(self.ent_update2.get())
            if self.ent_update3.get() != "":
                self.ent_index.append(self.ent_update3.get())
            if self.ent_update4.get() != "":
                self.ent_index.append(self.ent_update4.get())
            if self.ent_update5.get() != "":
                self.ent_index.append(self.ent_update5.get())
            if self.ent_update6.get() != "":
                self.ent_index.append(self.ent_update6.get())
            if self.ent_update7.get() != "":
                self.ent_index.append(self.ent_update7.get())
            if self.ent_update8.get() != "":
                self.ent_index.append(self.ent_update8.get())
            if self.ent_update9.get() != "":
                self.ent_index.append(self.ent_update9.get())
            if self.ent_update10.get() != "":
                self.ent_index.append(self.ent_update10.get())
            keyword = dt.datetime.strptime(self.keyword.get(), '%Y-%m-%d')
            data_index_use = self.df_sales_management[self.df_sales_management['日付'] == keyword]
            for (index, num) in zip(data_index_use.index, self.ent_index):
                self.df_sales_management.at[index, "販売個数"] = num
            
            workbook = openpyxl.load_workbook("../output/sales_management.xlsx")
            #ワークシート指定
            sheet = workbook['販売個数']

            sales_quantity = 18 # 販売個数のカラム番号
            date_column = 2
            yen_650 = 9 # 650円のカラム番号
            yen_550 = 10 # 550円のカラム番号
            yen_450 = 12 # 450円のカラム番号
            discount_value = 19 # 割引等単価のカラム番号
            coupon_quantity = 20 # クーポン利用のカラム番号
            menu_name_column = 6 # 弁当名等のカラム番号、50円引きクーポン利用を挿入するために使用

            for (index, menu_name, num) in zip(data_index_use.index, data_index_use['弁当名等'], self.ent_index):
                price = 0
                for value in (self.df_menu_master[self.df_menu_master['弁当名等']==menu_name])["販売価格"]:
                    price = value
                if price == 650:
                    sheet.cell(index + 2, yen_650, value=int(num))
                elif price == 550:
                    sheet.cell(index + 2, yen_550, value=int(num))
                elif price == 450:
                    sheet.cell(index + 2, yen_450, value=int(num))
            if self.ent_update_coupon.get() != "":
                row = (data_index_use.index[len(data_index_use.index)-1]+3).item()
                sheet.cell(row, menu_name_column, value='50円引きクーポン利用')
                sheet.cell(row, discount_value, value=-50)
                sheet.cell(row, coupon_quantity, value=int(self.ent_update_coupon.get()))
                sheet.cell(row, date_column, value=keyword)
            workbook.save('../output/sales_management.xlsx')

    def update_tree_by_search_result(self, result):
        self.tree.delete(*self.tree.get_children())
        self.result_text.set(f"検索結果：{len(result)}")
        for _, row in result.iterrows():
            self.tree.insert("", "end", values=row[self.colname_list].to_list())

    def onDouble(self, event):
        for item in self.tree.selection():
            print(self.tree.item(item)["values"])

def main():
    root = tk.Tk()
    app = MainWindow(master=root)
    app.mainloop()

if __name__ == "__main__":
    show_gui = main()